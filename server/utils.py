import os
import PyPDF2
from docx import Document
from docx.shared import Pt
import time
temp = int(time.time())
from ocrpdf import get_access_token,get_task_id,get_word_address,get_word_text

def read_pdf(file_path):
    print('111file_path:',file_path)
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        if text == '':
            access_token = get_access_token()
            task_id = get_task_id(file_path, access_token)
            word_downloadpath = get_word_address(task_id, access_token)
            print('word_downloadpath:',word_downloadpath)
            word_path = get_word_text(word_downloadpath,task_id)
            print('word_path:',word_path)
            return read_docx(word_path)
        else:
            return text

def read_docx(file_path):
    print('docfile_path:',file_path)
    full_text = []
    doc = Document(file_path)
    # 遍历文档中的所有段落
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 遍历文档中的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    print('\n'.join(full_text))
    # 将所有文本合并成一个字符串，并返回
    return '\n'.join(full_text)

def GetFileContent(file_path):
    print('file_path:',file_path)
    # 检查文件扩展名
    if file_path.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.endswith('.docx'):
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

# GetFileContent('./ocrpdftoword/Sq4rqSuU.docx')


def JsontoWord(directory, data):
    # 创建一个新的Word文档
    doc = Document()
    title = data.get('项目名称')
    # 添加标题
    doc.add_heading(f'{title}投标策划书', 0)

    # 创建表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # 设置表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '名称'
    hdr_cells[1].text = '内容'

    # 填充表格数据
    for key, value in data.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    # 自动调整表格列宽
    for cell in table.columns[0].cells:
        cell.width = Pt(100)  # 设置列宽，Pt(100)表示100磅

    # 保存Word文档
    path = os.path.join(directory, f"{title}投标策划书{temp}.docx")
    doc.save(path)
    return f"{title}投标策划书{temp}.docx"


